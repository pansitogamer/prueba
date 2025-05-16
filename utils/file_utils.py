import os
import shutil
from docx import Document
import PyPDF2

def guardar_archivo_docx(texto, ruta_salida):
    try:
        os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
        doc = Document()
        for linea in texto.strip().split('\n'):
            doc.add_paragraph(linea.strip())
        doc.save(ruta_salida)
        print(f"✅ Archivo guardado en: {ruta_salida}")
    except Exception as e:
        print(f"❌ Error al guardar archivo .docx: {e}")

def leer_docx(ruta_archivo):
    try:
        doc = Document(ruta_archivo)
        texto = "\n".join([p.text for p in doc.paragraphs])
        return texto.strip()
    except Exception as e:
        print(f"❌ Error al leer archivo .docx: {e}")
        return ""

def leer_pdf(ruta_archivo):
    try:
        texto = ""
        with open(ruta_archivo, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                texto += page.extract_text() + "\n"
        return texto.strip()
    except Exception as e:
        print(f"❌ Error al leer PDF: {e}")
        return ""

def cargar_prompt(ruta_prompt):
    try:
        with open(ruta_prompt, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"❌ Error al cargar prompt desde {ruta_prompt}: {e}")
        return ""

def guardar_archivos_candidato(nombre_completo, email, fecha_inicio, output_dir="output"):
    try:
        carpeta_candidato = os.path.join(output_dir, nombre_completo)
        os.makedirs(carpeta_candidato, exist_ok=True)

        
        return carpeta_candidato

    except Exception as e:
        print(f"❌ Error al guardar archivos para {nombre_completo}: {e}")
        return None
