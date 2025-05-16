import os
from dotenv import load_dotenv
from imap_tools import MailBox, AND
from builders.build_correo_credenciales import enviar_correo_credenciales
from docx import Document
import re
import json

# Cargar variables de entorno
load_dotenv()

IMAP_USER = os.getenv("IMAP_USER")
IMAP_PASSWORD = os.getenv("IMAP_PASSWORD")
IMAP_SERVER = os.getenv("IMAP_SERVER", "imap.gmail.com")
INPUT_DIR = "input"
OUTPUT_DIR = "output"

RESPUESTAS_POSITIVAS = ["sí", "si", "acepto", "quiero continuar", "confirmo"]

def respuesta_es_positiva(texto):
    texto = texto.lower()
    return any(r in texto for r in RESPUESTAS_POSITIVAS)

def comparar_emails(email1, email2):
    return email1.strip().lower() == email2.strip().lower()

def extraer_datos_metadatos(ruta_metadatos):
    nombre_completo = None
    rol_primario = None
    email = None
    try:
        doc = Document(ruta_metadatos)
        texto_total = ""

        for para in doc.paragraphs:
            texto_total += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto_total += cell.text + "\n"

        match_nombre = re.search(r"(?i)Nombre completo:\s*([^\n,]+)", texto_total)
        if match_nombre:
            nombre_completo = match_nombre.group(1).strip()
            print(f"✅ Nombre: {nombre_completo}")

        match_rol = re.search(r"(?i)Rol primario:\s*([^\n,]+)", texto_total)
        if match_rol:
            rol_primario = match_rol.group(1).strip()
            print(f"🎯 Rol primario: {rol_primario}")

        match_email = re.search(r"(?i)(Email|Correo):\s*([\w\.-]+@[\w\.-]+)", texto_total)
        if match_email:
            email = match_email.group(2).strip().lower()
            print(f"📧 Email extraído: {email}")

        return nombre_completo, rol_primario, email
    except Exception as e:
        print(f"❌ Error al leer metadatos: {e}")
        return None, None, None

def obtener_info_desde_links_json(nombre_completo):
    carpeta_output_candidato = os.path.join(OUTPUT_DIR, nombre_completo)
    path_links = os.path.join(carpeta_output_candidato, "links.json")

    if not os.path.exists(path_links):
        print(f"⚠️ No se encontró links.json para {nombre_completo}")
        return None, None

    with open(path_links, "r", encoding="utf-8") as f:
        data = json.load(f)
        link_credenciales = data.get("credenciales")
        link_responsabilidades = data.get("responsabilidades")
        return link_credenciales, link_responsabilidades

def buscar_y_procesar_respuestas():
    print("📬 Buscando respuestas de candidatos...")
    with MailBox(IMAP_SERVER).login(IMAP_USER, IMAP_PASSWORD, initial_folder="INBOX") as mailbox:
        mensajes = mailbox.fetch(AND(seen=False))

        for msg in mensajes:
            print(f"📧 Correo recibido: {msg.from_} - Asunto: {msg.subject}")
            remitente = msg.from_.lower().strip()
            contenido = msg.text.strip()

            if respuesta_es_positiva(contenido):
                print(f"✅ Respuesta positiva detectada de: {remitente}")
                candidato_encontrado = False
                nombre_completo = None
                rol_primario = None
                carpeta_output_candidato = None

                for folder in os.listdir(INPUT_DIR):
                    folder_path = os.path.join(INPUT_DIR, folder)
                    if not os.path.isdir(folder_path):
                        continue

                    for file in os.listdir(folder_path):
                        if file.endswith(".docx"):
                            ruta_metadatos = os.path.join(folder_path, file)
                            nombre, rol, email_extraido = extraer_datos_metadatos(ruta_metadatos)
                            print(f"🔍 Buscando coincidencia en archivo: {file}")
                            print(f"   ➤ Extraído: {email_extraido} | Remitente: {remitente}")

                            if email_extraido and comparar_emails(email_extraido, remitente):
                                print(f"✅ Coincidencia encontrada: {nombre} ({email_extraido})")
                                candidato_encontrado = True
                                nombre_completo = nombre
                                rol_primario = rol
                                carpeta_output_candidato = os.path.join(OUTPUT_DIR, nombre_completo)
                                break
                    if candidato_encontrado:
                        break

                if not candidato_encontrado:
                    print(f"⚠️ No se encontró un candidato cuyo email coincida con {remitente}.")
                    continue

                # Obtener los enlaces de credenciales y responsabilidades desde el archivo links.json
                link_credenciales, link_responsabilidades = obtener_info_desde_links_json(nombre_completo)

                if not link_credenciales or not link_responsabilidades:
                    print("⚠️ No se pudieron recuperar los enlaces desde links.json.")
                    continue

                decision = input(f"📤 ¿Deseas enviar correo de credenciales y roles a {nombre_completo}? (s/n): ").strip().lower()

                if decision == "s":
                    fecha_inicio = "Por definir"
                    # Enviar el correo de bienvenida solo con los enlaces, sin archivos adjuntos
                    enviar_correo_credenciales(nombre_completo, remitente, fecha_inicio, link_credenciales, link_responsabilidades, rol_primario)
                    print(f"✅ Correo de bienvenida con los enlaces enviados a {remitente}. Rol: {rol_primario}")
                else:
                    print("⏭️ Saltando envío para este candidato.")
            else:
                print(f"❌ Respuesta no interpretada como aceptación de: {remitente}")

if __name__ == "__main__":
    buscar_y_procesar_respuestas()
